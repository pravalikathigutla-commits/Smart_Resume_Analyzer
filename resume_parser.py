import pdfplumber
from docx import Document


# ============================================================
# PDF TEXT EXTRACTION
# ============================================================

def extract_from_pdf(file):
    """
    Extract text from a PDF file.
    """

    text_parts = []

    with pdfplumber.open(file) as pdf:

        for page in pdf.pages:

            page_text = page.extract_text()

            if page_text:
                text_parts.append(page_text)

    return "\n".join(text_parts)


# ============================================================
# DOCX TEXT EXTRACTION
# ============================================================

def extract_from_docx(file):
    """
    Extract text from a DOCX file.
    """

    document = Document(file)

    text_parts = []

    # Extract normal paragraphs
    for paragraph in document.paragraphs:

        paragraph_text = paragraph.text.strip()

        if paragraph_text:
            text_parts.append(paragraph_text)

    # Extract text from tables
    for table in document.tables:

        for row in table.rows:

            row_text = []

            for cell in row.cells:

                cell_text = cell.text.strip()

                if cell_text:
                    row_text.append(cell_text)

            if row_text:
                text_parts.append(" | ".join(row_text))

    return "\n".join(text_parts)


# ============================================================
# MAIN RESUME TEXT EXTRACTION
# ============================================================

def extract_resume_text(file, filename):
    """
    Extract text from PDF or DOCX resume.

    Parameters
    ----------
    file:
        Uploaded file object.

    filename:
        Name of the uploaded file.

    Returns
    -------
    str
        Extracted resume text.
    """

    if not filename:
        raise ValueError("Filename is required.")

    filename = filename.lower().strip()

    if filename.endswith(".pdf"):

        return extract_from_pdf(file)

    elif filename.endswith(".docx"):

        return extract_from_docx(file)

    else:

        raise ValueError(
            "Unsupported file type. Please upload PDF or DOCX."
        )